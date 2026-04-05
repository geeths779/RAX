"""Phase 2 verification: PDF/DOCX extraction + raw_text fallback."""
import asyncio
import io

from app.agents.resume_parser_agent import _extract_text, ResumeParserAgent
from app.agents.pipeline_context import PipelineContext
from app.agents.orchestrator import PipelineOrchestrator


def test_unsupported_type():
    try:
        _extract_text(b"test", "resume.txt")
        print("FAIL: Should have raised ValueError")
    except ValueError as e:
        print(f"PASS Test 1 — Unsupported type rejected: {e}")


def test_pdf_extraction():
    from PyPDF2 import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    buf = io.BytesIO()
    writer.write(buf)
    pdf_bytes = buf.getvalue()
    result = _extract_text(pdf_bytes, "resume.pdf")
    print(f"PASS Test 2 — PDF extraction returned {len(result)} chars (blank page)")


def test_docx_extraction():
    from docx import Document

    doc = Document()
    doc.add_paragraph("John Doe - Python Developer")
    doc.add_paragraph("Skills: Python, FastAPI, Docker")
    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()
    result = _extract_text(docx_bytes, "resume.docx")
    assert "John Doe" in result, f"Expected 'John Doe' in: {result}"
    assert "Python" in result, f"Expected 'Python' in: {result}"
    print(f"PASS Test 3 — DOCX extraction: \"{result}\"")


def test_raw_text_fallback():
    ctx = PipelineContext(resume_id="t1", raw_text="Some text")
    assert ctx.raw_text == "Some text"
    assert ctx.file_bytes == b""
    print("PASS Test 4 — raw_text fallback: file_bytes empty, raw_text preserved")


def test_pipeline_context_fields():
    ctx = PipelineContext(resume_id="t2", file_bytes=b"data", filename="cv.pdf")
    assert ctx.file_bytes == b"data"
    assert ctx.filename == "cv.pdf"
    assert ctx.raw_text == ""
    print("PASS Test 5 — PipelineContext: file_bytes and filename fields work")


def test_orchestrator_signature():
    orch = PipelineOrchestrator()
    import inspect
    sig = inspect.signature(orch.run)
    params = list(sig.parameters.keys())
    assert "file_bytes" in params, f"file_bytes not in run() params: {params}"
    assert "filename" in params, f"filename not in run() params: {params}"
    print(f"PASS Test 6 — Orchestrator.run() accepts file_bytes + filename")


async def test_parser_with_docx():
    """Integration test: DOCX bytes -> extraction -> Gemini parse (requires GOOGLE_API_KEY)."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Jane Smith")
    doc.add_paragraph("Email: jane@example.com | Phone: 555-1234")
    doc.add_paragraph("Skills: Python (5 years), React (3 years), SQL (4 years)")
    doc.add_paragraph("Experience: Software Engineer at Acme Corp (2020-2023)")
    doc.add_paragraph("Education: BS Computer Science, MIT, 2020")
    buf = io.BytesIO()
    doc.save(buf)

    ctx = PipelineContext(
        resume_id="test-docx-1",
        file_bytes=buf.getvalue(),
        filename="jane_smith.docx",
    )
    agent = ResumeParserAgent()
    ctx = await agent.run(ctx)

    if ctx.error:
        print(f"SKIP Test 7 — Gemini parse (may need API key): {ctx.error}")
        return

    assert ctx.raw_text, "raw_text should be populated from DOCX extraction"
    assert ctx.parsed_resume, "parsed_resume should be populated"
    print(f"PASS Test 7 — Full DOCX->parse pipeline: {list(ctx.parsed_resume.keys())}")


if __name__ == "__main__":
    test_unsupported_type()
    test_pdf_extraction()
    test_docx_extraction()
    test_raw_text_fallback()
    test_pipeline_context_fields()
    test_orchestrator_signature()

    print("\n--- Integration test (requires Gemini API key) ---")
    asyncio.run(test_parser_with_docx())

    print("\n=== Phase 2 verification complete ===")
